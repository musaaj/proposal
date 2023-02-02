#!/usr/bin/python3
import openai
from openai.error import ServiceUnavailableError, APIConnectionError
import sys
import os
import docx
import shutil
from docx import Document
from docx.opc.exceptions import OpcError, PackageNotFoundError
import ai
import tkinter
from tkinter.filedialog import askdirectory
from tkinter import Tk
from tkinter import ttk
from tkinter import filedialog
openai.api_key = ai.load_key('{}/key.json'.format(sys.path[0]))


class App(Tk):
    def __init__(self):
        super().__init__()
        self.title('GPTChat')
        self.geometry('600x520')
        self.prompt_var = tkinter.StringVar(self, value='Type here')
        self.directory_var = tkinter.StringVar(self, value='Choose Directory')
        self.progress_var = tkinter.StringVar(self, value='')
        self.create()

    def create(self):
        self.frame = tkinter.Frame(self, padx=10, pady=0)
        self.directory = ttk.Button(self.frame
                    ,textvariable=self.directory_var
                    ,command=self.on_directory
                )
        self.prompt = tkinter.Text(self.frame, padx=10, pady=10)
        self.progress = ttk.Label(self.frame, textvariable=self.progress_var)
        self.submit = ttk.Button(self.frame
                    ,text='OKAY'
                    ,width=500
                    ,command=self.on_submit
                )
        self.directory.pack()
        self.prompt.pack()
        self.progress.pack()
        self.submit.pack()
        self.frame.pack()

    def on_directory(self):
        folder = askdirectory(initialdir=self.directory_var.get());
        self.directory_var.set(folder)


    def on_submit(self):
        directory = self.directory_var.get()
        prompt = self.prompt.get('1.0', 'end')
        files = [file for file in os.listdir(directory) if file[-5:] == '.docx']
        self.backup(directory, files)
        self.progress_var.set('Contacting server...')
        for file in files:
            try:
                doc = Document(directory + '/' + file)
                self.process(doc, file, prompt)
            except PackageNotFoundError:
                self.progress_var.set('Error: invalid docx file')
                break

    def backup(self, directory, files):
        length = len(files)
        i = 0
        self.progress_var.set('Creating backup...')
        if not os.path.isdir(directory + '/backup'):
            os.mkdir(directory + '/backup')
        for file in files:
            i += 1
            self.progress_var.set('{}%'.format(int(i / length * 100)))
            if not os.path.isfile(directory + '/backup/' + file):
                shutil.copy(directory + '/' + file, directory + '/backup/' + file)
        self.progress_var.set('done')

    def process(self, doc, name, prompt):
        if not isinstance(doc, docx.document.Document):
            raise TypeError('doc must be a valid document object')
        if not isinstance(name, str):
            raise TypeError('name must be a string')
        if not isinstance(prompt, str):
            raise TypeError('prompt must be a string')
        length = len(doc.paragraphs)
        percent = 0
        self.progress_var.set('Working {}...0%'.format(name))
        for paragraph in doc.paragraphs:
            percent += 1
            self.progress_var.set(('Working on {}...{}%'.format(name
                    ,int((percent / length) * 100)))
                )
            if len(paragraph.text):
                try:
                    text = ai.get_result(prompt + paragraph.text)
                    paragraph.text = text
                except ServiceUnavailableError:
                    self.progress_var.set(('Error: Server overloaded!'))
                except APIConnectionError:
                    self.progress_var.set('Error: No internet connection')
        doc.save(self.directory_var.get() + '/' + name)




if __name__ == '__main__':
    App().mainloop()
